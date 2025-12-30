# -*- coding: utf-8 -*-
"""
将MySQL数据库表结构生成Word文档工具

该脚本用于连接MySQL数据库，读取指定数据库中的表结构信息，
并将这些信息导出到Word文档中，生成表结构说明文档。

功能特性：
1. 连接MySQL数据库并获取表结构信息
2. 生成包含表清单和各表详细字段信息的Word文档
3. 支持多个数据库的批量处理
4. 生成的文档包含表名、字段名、数据类型、约束和字段说明等信息
5. 文档格式美观，标题和表格清晰

使用方法：
1. 配置数据库连接信息
2. 设置需要导出的数据库名称列表
3. 运行脚本自动生成Word文档

依赖包：
- pymysql: 用于连接MySQL数据库
- python-docx: 用于生成Word文档

注意事项：
- 需要安装pymysql和python-docx包
- 需要正确配置数据库连接参数
"""

# 导包
# pip  install pymysql python-docx
import pymysql
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
 
# 读取或者创建文件
document = Document()       # 新建文档
# document = Document('test.docx')      # 读取现有文档，建立文档对象

# 建立数据库连接
<<<<<<< HEAD
conn = pymysql.connect(host="192.168.100.21",
                           port=33306,
                           user="root",
                           password="123456",
                        #    db="db_mingyang_home_court",
=======
conn = pymysql.connect(host="127.0.0.1",
                           port=3306,
                           user="root",
                           password="123456",
                        #    db="db_test_home_court",
>>>>>>> 3c4181f7e07feb1892a8b680b63799f5f5dfde4c
                           charset="utf8")
cursor = conn.cursor()

def tb_list(db_name):
    '''
    查询指定数据库中的所有数据表
    参数:
        db_name: 数据库名称
    返回:
        包含表名和表注释的元组列表
    '''
    sql='''
        SELECT
            TABLE_NAME,
            TABLE_COMMENT 
        FROM
            information_schema.TABLES 
        WHERE
            table_schema = '{0}' 
            AND TABLE_COMMENT NOT LIKE '%弃用%';
    '''.format(db_name)

    print('1，执行sql：',sql)
    cursor.execute(sql)
    result = cursor.fetchall()
    return result


def load_data_from_mysql(db_name, table_name):
    '''
    从MySQL数据库中加载指定表的字段信息
    参数:
        db_name: 数据库名称
        table_name: 表名
    返回:
        包含字段信息的元组列表（字段名、逻辑名、数据类型、约束、说明）
    '''
    sql='''
    SELECT 
        t.COLUMN_NAME 字段名,
        t.COLUMN_COMMENT 逻辑名,
        IF(LEN IS NOT NULL,CONCAT(DATA_TYPE,'(',LEN,')'),DATA_TYPE) 数据类型,
        CASE
                WHEN t.column_key = 'PRI' 
                THEN 'PK' 
                ELSE '' 
            END 约束,
        t.COLUMN_COMMENT 说明 
    FROM (
        SELECT 
            COLUMN_NAME,
            COLUMN_COMMENT,
            DATA_TYPE,
            (
                CASE
                    WHEN DATA_TYPE = 'float' 
                    OR DATA_TYPE = 'double' 
                    OR DATA_TYPE = 'tinyint' 
                    OR DATA_TYPE = 'SMALLINT' 
                    OR DATA_TYPE = 'MEDIUMINT' 
                    OR DATA_TYPE = 'INT' 
                    OR DATA_TYPE = 'INTEGER' 
                    OR DATA_TYPE = 'decimal' 
                    OR DATA_TYPE = 'bigint' 
                    THEN NUMERIC_PRECISION 
                    ELSE CHARACTER_MAXIMUM_LENGTH
                END
            ) + '' LEN,
        --   CASE
        --     WHEN IS_NULLABLE = 'YES' 
        --     THEN '是' 
        --     ELSE '否' 
        --   END 空否,
            column_key
        FROM
            INFORMATION_SCHEMA.COLUMNS 
        WHERE table_schema = '{0}' 
            AND table_name = '{1}'
    ) AS t
    '''.format(db_name, table_name)

    print('2，执行sql：',sql)
    cursor.execute(sql)
    result = cursor.fetchall()
    return result
 

def set_heading(content1, level, content2):
    '''
    设置文档标题样式
    参数:
        content1: 标题前缀内容
        level: 标题级别（1-9）
        content2: 标题后缀内容
    '''
    # 第二种设置标题的方式，此方式还可以设置文本的字体、颜色、大小等属性
    run = document.add_heading(content1, level).add_run(content2)
    # 设置西文字体
    run.font.name = u'宋体'
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 设置字体颜色
    run.font.color.rgb = RGBColor(0,0,0)  # 黑色
    # 设置字体大小
    # run.font.size = Pt(30)
    # 设置下划线
    # run.font.underline = True
    # 设置删除线
    # run.font.strike = True
    # 设置加粗
    run.bold = True
    # 设置斜体
    # run.italic = True
 


# 总的大标题部分
set_heading('', 1, '表结构')


'''
要查询的表初始化
'''
# 需要查询的库名
<<<<<<< HEAD
# db_name_list = ['db_mingyang_builder']
db_name_list = ['db_mingyang_builder', 'db_mingyang_certificate', 'db_mingyang_common', 'db_mingyang_common_service', 'db_mingyang_home_court', 'db_mingyang_matching', 'db_mingyang_pay', 'db_mingyang_pubilc', 'db_mingyang_venue_booking', 'db_mingyang_venue_services_manage', 'db_seata']
=======
# db_name_list = ['db_test_builder']
db_name_list = ['db_test_builder', 'db_test_certificate', 'db_test_common', 'db_test_common_service', 'db_test_home_court', 'db_test_matching', 'db_test_pay', 'db_test_pubilc', 'db_test_venue_booking', 'db_test_venue_services_manage', 'db_seata']
>>>>>>> 3c4181f7e07feb1892a8b680b63799f5f5dfde4c

 
'''
写入小标题以及表格内容
'''
# 设置正文全局字体
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
 
# for i in range(0, len(db_name_list)):
#     print (i, db_name_list[i])


for i in range(len(db_name_list)):
    # 库名
    db_name = db_name_list[i]
    # print('db_name_list---', db_name_list[i])
    tbs = tb_list(db_name)

    # 写入小标题
    set_heading('', 5, '表清单')
    # document.add_heading('表清单', level=3)
    # 表格标题
    table_name = ['表名', '功能说明']
    # 二维元组转化为二维列表 ,且将其他类型转化为字符串 且将None转化为 ''
    table_data_list = list(list([( '' if it is None else str(it)) for it in items]) for items in list(tbs))
    # 创建表格行列
    table = document.add_table(rows=len(tbs)+1, cols=len(table_name),style='Table Grid')
    # 首行设置背景色
    rows = table.rows[0]
    for cell in rows.cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    # 写入表格标题
    for i in range(len(table_name)):
        cell = table.cell(0, i)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中
        cell.paragraphs[0].add_run(table_name[i]).font.bold = True   # 加粗 
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中

    # 写入表格内容
    for i in range(len(table_data_list)):
        for j in range(len(table_name)):
            table.cell(i+1, j).text = table_data_list[i][j]
    


    for tb in tbs:
        print('tb：',tb[0], tb[1])
        table_name = tb[0]
        table_name_desc = tb[1]

        # 查询sql获得二维元组
        tuple1 = load_data_from_mysql(db_name, table_name)
        # 二维元组转化为二维列表 ,且将其他类型转化为字符串 且将None转化为 ''
        table_data_list = list(list([( '' if it is None else str(it)) for it in items]) for items in list(tuple1))
        
        # 写入小标题
        title = table_name + ' ( ' + table_name_desc + '表 )'
        # document.add_heading(title, level=5)
        set_heading('', 6, title)
        # 表格标题
        table_name = ['字段名', '逻辑名', '数据类型', '约束', '说明']
        # 创建表格行列
        table = document.add_table(rows=len(table_data_list)+1, cols=len(table_name),style='Table Grid')
        # 首行设置背景色
        rows = table.rows[0]
        for cell in rows.cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        # 写入表格标题
        for i in range(len(table_name)):
            cell = table.cell(0, i)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中
            cell.paragraphs[0].add_run(table_name[i]).font.bold = True   # 加粗 
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中
    
        # 写入表格内容
        for i in range(len(table_data_list)):
            for j in range(len(table_name)):
                table.cell(i+1, j).text = table_data_list[i][j]
 
 
# 保存文档
document.save("test.docx")